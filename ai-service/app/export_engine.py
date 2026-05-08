"""PDF (reportlab) and DOCX export from draft JSON."""
from io import BytesIO
from typing import Any

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document


def _ordered_sections(draft: dict[str, Any]) -> list[tuple[str, Any]]:
    order = draft.get("sectionOrder") or []
    sections = draft.get("sections") or {}
    out = []
    for key in order:
        if key in sections:
            out.append((key, sections[key]))
    for key, val in sections.items():
        if key not in order:
            out.append((key, val))
    return out


def export_pdf(draft: dict[str, Any]) -> bytes:
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    personal = (draft.get("sections") or {}).get("personal") or {}
    title = personal.get("fullName") or "Resume"
    c.drawString(50, y, title[:80])
    y -= 22
    c.setFont("Helvetica", 10)
    line = " | ".join(
        filter(
            None,
            [
                personal.get("email", ""),
                personal.get("phone", ""),
                personal.get("location", ""),
            ],
        )
    )
    if line:
        c.drawString(50, y, line[:100])
        y -= 16
    if personal.get("summary"):
        c.setFont("Helvetica-Oblique", 10)
        for piece in _wrap(personal["summary"], 95):
            c.drawString(50, y, piece)
            y -= 14
            if y < 72:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica-Oblique", 10)
        y -= 8
    c.setFont("Helvetica", 10)
    for section_key, val in _ordered_sections(draft):
        if section_key == "personal":
            continue
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, section_key.upper())
        y -= 14
        c.setFont("Helvetica", 10)
        lines = _section_lines(section_key, val)
        for piece in lines:
            for wline in _wrap(piece, 95):
                c.drawString(50, y, wline)
                y -= 12
                if y < 72:
                    c.showPage()
                    y = height - 50
                    c.setFont("Helvetica", 10)
        y -= 6
    c.save()
    return buf.getvalue()


def _wrap(text: str, width: int) -> list[str]:
    words = text.split()
    if not words:
        return []
    lines = []
    cur = []
    for w in words:
        if sum(len(x) + 1 for x in cur) + len(w) <= width:
            cur.append(w)
        else:
            lines.append(" ".join(cur))
            cur = [w]
    if cur:
        lines.append(" ".join(cur))
    return lines


def _section_lines(key: str, val: Any) -> list[str]:
    if key == "skills" and isinstance(val, list):
        return [", ".join(val)]
    if key == "experience" and isinstance(val, list):
        lines = []
        for item in val:
            if not isinstance(item, dict):
                lines.append(str(item))
                continue
            head = " — ".join(
                filter(None, [item.get("role") or item.get("degree"), item.get("company") or item.get("school")])
            )
            if head:
                lines.append(head)
            if item.get("description"):
                lines.append(item["description"])
            if item.get("field"):
                lines.append(item["field"])
            if item.get("start") or item.get("end"):
                lines.append(" — ".join(filter(None, [item.get("start"), item.get("end")])))
        return lines
    if key == "education" and isinstance(val, list):
        lines = []
        for item in val:
            if not isinstance(item, dict):
                lines.append(str(item))
                continue
            head = " — ".join(filter(None, [item.get("degree"), item.get("school")]))
            if head:
                lines.append(head)
            if item.get("field"):
                lines.append(item["field"])
            if item.get("end"):
                lines.append(str(item["end"]))
        return lines
    if key == "projects" and isinstance(val, list):
        lines = []
        for item in val:
            if not isinstance(item, dict):
                lines.append(str(item))
                continue
            title = item.get("title") or item.get("name") or ""
            row = title
            ts = item.get("techStack") or item.get("tech_stack")
            if isinstance(ts, list) and ts:
                row = f"{row} — Tech: {', '.join(str(x) for x in ts)}" if title else f"Tech: {', '.join(str(x) for x in ts)}"
            links = []
            if item.get("github"):
                links.append(f"GitHub: {item['github']}")
            lk = item.get("liveLink") or item.get("live_link")
            if lk:
                links.append(f"Live: {lk}")
            if title or links:
                lines.append(" · ".join([x for x in [row if title else None, *links] if x]))
            if item.get("description"):
                lines.append(item["description"])
        return lines
    if key == "certifications" and isinstance(val, list):
        lines = []
        for item in val:
            if isinstance(item, dict):
                lines.append(
                    " — ".join(
                        filter(
                            None,
                            [
                                item.get("name"),
                                item.get("issuer"),
                                str(item.get("issueDate")) if item.get("issueDate") else None,
                            ],
                        )
                    )
                )
            else:
                lines.append(str(item))
        return [x for x in lines if x]
    if key == "achievements" and isinstance(val, list):
        lines = []
        for item in val:
            if isinstance(item, dict):
                t = item.get("title")
                if t:
                    lines.append(str(t))
                if item.get("description"):
                    lines.append(str(item["description"]))
            else:
                lines.append(str(item))
        return lines
    if key == "languages" and isinstance(val, list):
        lines = []
        for item in val:
            if isinstance(item, dict):
                lines.append(
                    " — ".join(filter(None, [item.get("language") or item.get("name"), item.get("proficiency")]))
                )
            else:
                lines.append(str(item))
        return [x for x in lines if x]
    if isinstance(val, dict):
        return [f"{k}: {v}" for k, v in val.items() if v]
    if val:
        return [str(val)]
    return []


def export_docx(draft: dict[str, Any]) -> bytes:
    doc = Document()
    personal = (draft.get("sections") or {}).get("personal") or {}
    doc.add_heading(personal.get("fullName") or "Resume", level=0)
    sub = " | ".join(
        filter(None, [personal.get("email"), personal.get("phone"), personal.get("location")])
    )
    if sub:
        doc.add_paragraph(sub)
    if personal.get("summary"):
        doc.add_paragraph(personal["summary"])
    for section_key, val in _ordered_sections(draft):
        if section_key == "personal":
            continue
        doc.add_heading(section_key.capitalize(), level=1)
        for line in _section_lines(section_key, val):
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
